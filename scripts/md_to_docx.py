"""Convert a Markdown note (with headings, tables, lists, bold/inline-code) to a .docx file."""
import sys
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_runs(paragraph, text):
    """Add text to paragraph, handling **bold** and `code` inline markers."""
    pattern = re.compile(r'(\*\*.+?\*\*|`.+?`)')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
        else:
            paragraph.add_run(part)


def parse_table(lines, start):
    """Parse a markdown table starting at lines[start]. Returns (rows, next_index)."""
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith('|'):
        line = lines[i].strip()
        if re.match(r'^\|[\s:|-]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i


def convert(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    if rfonts is None:
        from docx.oxml.ns import qn
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    from docx.oxml.ns import qn
    rfonts.set(qn('w:eastAsia'), '微软雅黑')

    i = 0
    in_code_block = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == '---':
            doc.add_paragraph('—' * 20)
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            heading_level = min(level, 9)
            p = doc.add_heading(level=heading_level)
            add_runs(p, text)
            i += 1
            continue

        # Tables
        if stripped.startswith('|'):
            rows, i = parse_table(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Table Grid'
                for r, row in enumerate(rows):
                    for c, cell_text in enumerate(row):
                        if c < len(table.rows[r].cells):
                            cell = table.rows[r].cells[c]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            add_runs(p, cell_text)
                            if r == 0:
                                for run in p.runs:
                                    run.bold = True
            continue

        # Blockquote
        if stripped.startswith('>'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_runs(p, stripped.lstrip('> ').strip())
            i += 1
            continue

        # Bullet list
        m = re.match(r'^(\s*)-\s+(.*)$', line)
        if m:
            indent = len(m.group(1))
            text = m.group(2)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Pt(18 + indent * 9)
            add_runs(p, text)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\s*)(\d+)\.\s+(.*)$', line)
        if m:
            text = m.group(3)
            p = doc.add_paragraph(style='List Number')
            add_runs(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f'Saved: {docx_path}')


if __name__ == '__main__':
    convert(sys.argv[1], sys.argv[2])
